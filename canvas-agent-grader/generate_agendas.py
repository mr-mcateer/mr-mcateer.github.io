#!/usr/bin/env python3
"""Generate Monday April 13, 2026 agendas as Word docs."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

DOWNLOADS = os.path.expanduser("~/Downloads")


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    return h


def add_body(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    run.bold = bold
    run.italic = italic
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = "Calibri"
    return p


def build_autos_agenda():
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Title
    title = doc.add_heading("AUTOMOTIVE MANUFACTURING AGENDA", level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    add_body(doc, "Monday, April 13, 2026  --  ALL Day (50 minutes)", bold=True)
    add_body(doc, "Mr. McAteer  --  P1 Engines/Fab")

    # Week at a glance
    add_heading_styled(doc, "WEEK AT A GLANCE", level=2)
    add_bullet(doc, "Monday 4/13: ALL day (today) -- 50 min | Staff Meeting 3:20 PM")
    add_bullet(doc, "Tuesday 4/14: Early Release ODD (2:15 dismissal) -- 60 min")
    add_bullet(doc, "Wednesday 4/15: Early Release Even (12:50 dismissal) -- no shop classes")
    add_bullet(doc, "Thursday 4/16: Drop-In Conferences (no school for students)")
    add_bullet(doc, "Friday 4/17: No school")
    add_body(doc, "This is your last full-length class before conferences. Use it well.", italic=True)

    # Daily Journal
    add_heading_styled(doc, "DAILY JOURNAL (5 min)", level=2)
    add_body(doc, "You just did a multi-point vehicle inspection on Thursday. What is one thing you checked that you had never thought about before? If you were buying a used car this weekend, what would be the first three things you inspect and why?")

    # Priorities
    add_heading_styled(doc, "TODAY'S PRIORITIES", level=2)

    add_body(doc, "1. Canvas Module Work (35 min)", bold=True)
    add_body(doc, "Open Canvas and check your module progress. Work through the next uncompleted assignment at your own pace:")
    add_bullet(doc, "02 -- Everything You Should Check Before You Buy (25 pts): Record sticker data, decode your VIN, and compare factory vs. dealer options.")
    add_bullet(doc, "03 -- The Real Price Tag (40 pts): Calculate monthly payments, total interest, and total cost at 48, 60, and 72 months using 5.5% APR. Choose your preferred term and explain why.")
    add_bullet(doc, "04 -- Insurance Decoded (25 pts): Research liability, collision, and comprehensive coverage for your vehicle.")
    add_bullet(doc, "05 -- Keep It Running (30 pts): Build a maintenance schedule and estimate annual upkeep costs.")
    add_body(doc, "If you are caught up on modules, continue working on your Engine Teardown Portfolio. Photos and written assessments need to be in the document, not just the template.")

    add_body(doc, "2. Grade Check and Catch-Up (10 min)", bold=True)
    add_body(doc, "Before conferences, check your current grade in Canvas. If you have any missing or incomplete assignments, now is the time to fix them. I will be grading submitted work this week.")
    add_body(doc, "Students with pending resubmissions: check Canvas comments for feedback on what to fix.", italic=True)

    # Before you leave
    add_heading_styled(doc, "BEFORE YOU LEAVE", level=2)
    add_bullet(doc, "Submit whatever you have completed today, even if it is not perfect. Partial credit beats missing.")
    add_bullet(doc, "Log out of Chromebooks and return them to the cart.")
    add_bullet(doc, "Push in your chair and clean your area.")

    # Looking ahead
    add_heading_styled(doc, "LOOKING AHEAD", level=2)
    add_bullet(doc, "Tuesday 4/14 is early release (2:15 dismissal). ODD schedule gives you 60 minutes. Same work, same pace.")
    add_bullet(doc, "Wednesday 4/15 is early release even (12:50 dismissal). No shop classes.")
    add_bullet(doc, "Thursday: drop-in parent-teacher conferences. Grades will be updated before they begin.")
    add_bullet(doc, "When we return Monday 4/20, we continue with modules and start planning for the oil change rotation.")

    path = os.path.join(DOWNLOADS, "Autos_Agenda_04-13-2026.docx")
    doc.save(path)
    print(f"Saved: {path}")
    return path


def build_metals_agenda():
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Title
    title = doc.add_heading("METALS FABRICATION AGENDA", level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    add_body(doc, "Monday, April 13, 2026  --  ALL Day (45 minutes)", bold=True)
    add_body(doc, "Mr. McAteer  --  P3 Metals / P5 Metals")

    # Week at a glance
    add_heading_styled(doc, "WEEK AT A GLANCE", level=2)
    add_bullet(doc, "Monday 4/13: ALL day (today) -- 45 min | Staff Meeting 3:20 PM")
    add_bullet(doc, "Tuesday 4/14: Early Release ODD (2:15 dismissal) -- 60 min")
    add_bullet(doc, "Wednesday 4/15: Early Release Even (12:50 dismissal) -- no shop classes")
    add_bullet(doc, "Thursday 4/16: Drop-In Conferences (no school for students)")
    add_bullet(doc, "Friday 4/17: No school")
    add_body(doc, "This is your last full-length class before conferences. Make progress on your project today.", italic=True)

    # Daily Journal
    add_heading_styled(doc, "DAILY JOURNAL (5 min)", level=2)
    add_body(doc, "Name one tool or machine in the shop that you feel confident using now that you did not at the start of the semester. What changed?")

    # Priorities
    add_heading_styled(doc, "TODAY'S PRIORITIES", level=2)

    add_body(doc, "1. Open Shop / Project Work (30 min)", bold=True)
    add_body(doc, "Continue your current fabrication project. If you are between projects, pick up a Design Challenge or start a CSWA practice part.")
    add_bullet(doc, "Metal Bookend Project: Design, cut, weld. Bring completed work to Mr. McAteer for inspection.")
    add_bullet(doc, "Design Challenge: Tiles, scrap metal, spray paint, or Epilog laser. Get creative.")
    add_bullet(doc, "CSWA Practice Parts: Self-check your work against the answer key before moving on.")
    add_bullet(doc, "Plasma CNC: Use the calculator before cutting. Queue with Mr. McAteer if the machine is open.")
    add_bullet(doc, "Knee Mill: Check with Bob for rotation availability.")

    add_body(doc, "2. Grade Check and Missing Work (10 min)", bold=True)
    add_body(doc, "Conferences are Thursday. Check your Canvas grade now.")
    add_bullet(doc, "If you have 3 or more missing assignments, talk to Mr. McAteer today about a plan.")
    add_bullet(doc, "Shop Cleanup submissions, Knee Mill Quiz, and Shop Re-Entry Reflection are common gaps. Check your to-do list in Canvas.")
    add_bullet(doc, "If you owe a Metal Bookend writeup, submit what you have. Photo documentation counts.")

    # Safety + shop standards
    add_heading_styled(doc, "SHOP STANDARDS", level=2)
    add_bullet(doc, "Closed-toe shoes and safety glasses at all times on the shop floor.")
    add_bullet(doc, "Hearing protection at machines.")
    add_bullet(doc, "Blue tape with your name on all personal tools and projects.")
    add_bullet(doc, "Material prep: alcohol-wipe before paint, degrease before powder coat.")
    add_bullet(doc, "Report anything broken or unsafe immediately.")

    # Before you leave
    add_heading_styled(doc, "BEFORE YOU LEAVE", level=2)
    add_bullet(doc, "Return all tools to the correct drawer. Check the Master Tool List if unsure.")
    add_bullet(doc, "Sweep your station. Wipe down surfaces.")
    add_bullet(doc, "Submit any Canvas work you completed today.")
    add_bullet(doc, "Do not leave until Mr. McAteer gives the all-clear.")

    # Looking ahead
    add_heading_styled(doc, "LOOKING AHEAD", level=2)
    add_bullet(doc, "Tuesday 4/14 is early release (2:15 dismissal). ODD schedule gives you 60 minutes. Full shop access.")
    add_bullet(doc, "Wednesday 4/15 is early release even (12:50 dismissal). No shop classes.")
    add_bullet(doc, "Thursday: drop-in parent-teacher conferences. Grades will be updated before they begin.")
    add_bullet(doc, "When we return Monday 4/20, project work continues. Bookend and Design Challenge deadlines are approaching.")

    path = os.path.join(DOWNLOADS, "Metals_Agenda_04-13-2026.docx")
    doc.save(path)
    print(f"Saved: {path}")
    return path


if __name__ == "__main__":
    build_autos_agenda()
    build_metals_agenda()
    print("Done.")
