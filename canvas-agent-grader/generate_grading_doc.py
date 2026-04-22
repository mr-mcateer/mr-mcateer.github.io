#!/usr/bin/env python3
"""
Generate a one-stop grading reference document for all ungraded edge-case submissions.

Aggregates:
- Student info and current grade
- Assignment details and rubric criteria
- Full submission content
- Comment thread history
- SpeedGrader screenshot (embedded if available)
- Suggested grade with rationale
"""

import json
import os
import glob
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

DOWNLOADS = os.path.expanduser("~/Downloads")
EXPORTS = os.path.join(os.path.dirname(__file__), "exports")


def load_data():
    with open(os.path.join(EXPORTS, "detailed_audit.json")) as f:
        audit = json.load(f)
    with open(os.path.join(EXPORTS, "gradebook_snapshot.json")) as f:
        snapshot = json.load(f)
    return audit, snapshot


def get_student_score(snapshot, course_id, user_id):
    """Look up current score for a student in the snapshot."""
    cid = str(course_id)
    if cid not in snapshot.get("courses", {}):
        return None, None
    course = snapshot["courses"][cid]
    for student in course.get("enrollments", []):
        if str(student.get("user_id")) == str(user_id):
            return student.get("current_score"), student.get("final_score")
    return None, None


def suggest_grade(submission):
    """
    Suggest a grade based on submission content, type, and context.
    Returns (suggested_score, max_score, rationale).
    """
    stype = submission.get("submission_type")
    content = submission.get("content", [])
    comments = submission.get("comments", [])
    pts = submission.get("points_possible", 0)
    state = submission.get("workflow_state", "")
    aname = submission.get("assignment_name", "")

    content_text = "\n".join(content) if content else ""
    student_comments = [c for c in comments if c.get("author") != "Andy McAteer"]
    word_count = len(content_text.split()) if content_text else 0

    # No submission at all
    if not stype or content_text.strip() in ["", "(no submission)"]:
        if any("excused" in str(c.get("comment", "")).lower() for c in comments):
            return (None, pts, "Excused -- no action needed")
        return (0, pts, "No submission received. Consider 0 or excused.")

    # Deep Clean / Shop Cleanup -- participation credit
    if "clean" in aname.lower() or "cleanup" in aname.lower():
        if word_count > 10:
            return (pts, pts, "Participation-based. Student described cleanup work. Full credit recommended.")
        return (pts * 0.5, pts, "Minimal description of cleanup work. Consider partial credit.")

    # Text submissions for portfolio modules
    if stype == "online_text_entry":
        if word_count < 20:
            return (pts * 0.15, pts, f"Very sparse submission ({word_count} words). Minimal effort shown.")
        elif word_count < 50:
            return (pts * 0.4, pts, f"Partial submission ({word_count} words). Key sections likely incomplete.")
        elif word_count < 100:
            return (pts * 0.65, pts, f"Moderate submission ({word_count} words). Check rubric coverage.")
        else:
            return (pts * 0.85, pts, f"Substantial submission ({word_count} words). Review for rubric alignment.")

    # File uploads
    if stype == "online_upload":
        return (None, pts, "File upload -- needs manual review of document content.")

    # URL submissions (Google Docs)
    if stype == "online_url":
        return (None, pts, "URL submission -- needs manual review. Check document sharing permissions.")

    return (None, pts, "Could not auto-suggest. Manual review required.")


def add_section_heading(doc, text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    return p


def build_grading_doc():
    audit, snapshot = load_data()
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)

    # Title
    title = doc.add_heading("GRADING REFERENCE -- EDGE CASES", level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    p = doc.add_paragraph()
    run = p.add_run("Generated: April 12, 2026 | Pre-Conference Grading Sweep")
    run.font.size = Pt(11)
    run.bold = True

    p = doc.add_paragraph()
    run = p.add_run("20 submissions across 4 courses | 15 with content, 5 excused/no submission")
    run.font.size = Pt(10)
    run.italic = True

    doc.add_page_break()

    # Summary table
    add_section_heading(doc, "GRADING QUEUE SUMMARY")
    table = doc.add_table(rows=1, cols=6)
    table.style = "Light Grid Accent 1"
    headers = ["#", "Student", "Assignment", "Course", "Pts", "Status"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for p in table.rows[0].cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)

    count = 0
    all_submissions = []

    for cid, course in audit["courses"].items():
        cname = course.get("label", "")
        for assignment in course.get("assignments_needing_grading", []):
            aname = assignment.get("name", "")
            aid = assignment.get("id", "")
            pts = assignment.get("points_possible", 0)
            has_rubric = assignment.get("has_rubric", False)
            for sub in assignment.get("ungraded_submissions", []):
                count += 1
                sname = sub.get("student_name", "")
                uid = sub.get("user_id", "")
                stype = sub.get("submission_type")
                content = sub.get("content", [])
                comments = sub.get("comments", [])
                submitted = sub.get("submitted_at", "")
                late = sub.get("late", False)
                score_current, score_final = get_student_score(snapshot, cid, uid)

                status = "No submission" if not stype else ("Late" if late else "On time")

                row = table.add_row()
                row.cells[0].text = str(count)
                row.cells[1].text = sname
                row.cells[2].text = aname[:40]
                row.cells[3].text = cname.replace("P1 Engines Fab", "Fab").replace("P5 Metals", "P5M").replace("P3 Metals", "P3M")
                row.cells[4].text = str(int(pts))
                row.cells[5].text = status
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.size = Pt(8)

                # Enrich submission data for detailed section
                sub["assignment_name"] = aname
                sub["assignment_id"] = aid
                sub["points_possible"] = pts
                sub["has_rubric"] = has_rubric
                sub["course_name"] = cname
                sub["course_id"] = cid
                sub["current_score"] = score_current
                sub["final_score"] = score_final
                sub["entry_num"] = count
                all_submissions.append(sub)

    doc.add_page_break()

    # Detailed entries
    add_section_heading(doc, "DETAILED SUBMISSION REVIEWS")

    for sub in all_submissions:
        sname = sub.get("student_name", "")
        aname = sub.get("assignment_name", "")
        cname = sub.get("course_name", "")
        pts = sub.get("points_possible", 0)
        stype = sub.get("submission_type")
        content = sub.get("content", [])
        comments = sub.get("comments", [])
        submitted = sub.get("submitted_at", "")
        late = sub.get("late", False)
        score_current = sub.get("current_score")
        entry_num = sub.get("entry_num")
        cid = sub.get("course_id")
        aid = sub.get("assignment_id")
        uid = sub.get("user_id")

        # Entry header
        h = doc.add_heading(f"#{entry_num}: {sname} -- {aname}", level=3)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

        # Info table
        info_table = doc.add_table(rows=4, cols=2)
        info_table.style = "Light Grid Accent 1"
        info_data = [
            ("Course", cname),
            ("Current Grade", f"{score_current}%" if score_current else "N/A"),
            ("Points Possible", str(int(pts))),
            ("Submitted", f"{submitted or 'None'}{' (LATE)' if late else ''}"),
        ]
        for i, (label, value) in enumerate(info_data):
            info_table.rows[i].cells[0].text = label
            info_table.rows[i].cells[1].text = value
            for cell in info_table.rows[i].cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(9)

        doc.add_paragraph()

        # Submission content
        p = doc.add_paragraph()
        run = p.add_run("SUBMISSION CONTENT:")
        run.bold = True
        run.font.size = Pt(10)

        if content:
            for line in content:
                p = doc.add_paragraph()
                run = p.add_run(line[:2000])  # Truncate very long content
                run.font.size = Pt(9)
                run.font.name = "Consolas"
                p.paragraph_format.left_indent = Inches(0.25)
        else:
            p = doc.add_paragraph()
            run = p.add_run("(No submission content)")
            run.font.size = Pt(9)
            run.italic = True

        # Comment thread
        if comments:
            p = doc.add_paragraph()
            run = p.add_run(f"COMMENT THREAD ({len(comments)} messages):")
            run.bold = True
            run.font.size = Pt(10)

            for c in comments:
                author = c.get("author", "Unknown")
                date = c.get("created_at", "")[:10]
                text = c.get("comment", "")
                is_teacher = author == "Andy McAteer"

                p = doc.add_paragraph()
                run = p.add_run(f"[{date}] {author}: ")
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x00, 0x50, 0xA0) if is_teacher else RGBColor(0x33, 0x33, 0x33)

                run = p.add_run(text[:500])
                run.font.size = Pt(9)
                p.paragraph_format.left_indent = Inches(0.25)

        # SpeedGrader link
        sg_url = f"https://csd509j.instructure.com/courses/{cid}/gradebook/speed_grader?assignment_id={aid}&student_id={uid}"
        p = doc.add_paragraph()
        run = p.add_run("SpeedGrader: ")
        run.bold = True
        run.font.size = Pt(9)
        run = p.add_run(sg_url)
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x00, 0x50, 0xA0)

        # Suggested grade
        suggested, max_pts, rationale = suggest_grade(sub)

        p = doc.add_paragraph()
        run = p.add_run("SUGGESTED GRADE: ")
        run.bold = True
        run.font.size = Pt(11)

        if suggested is not None:
            score_text = f"{suggested:.0f} / {max_pts:.0f}"
            run = p.add_run(score_text)
            run.bold = True
            run.font.size = Pt(11)
            if suggested == 0:
                run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
            elif suggested >= max_pts * 0.8:
                run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
            else:
                run.font.color.rgb = RGBColor(0xCC, 0x88, 0x00)
        else:
            run = p.add_run("MANUAL REVIEW NEEDED")
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x00, 0x50, 0xA0)

        p = doc.add_paragraph()
        run = p.add_run(f"Rationale: {rationale}")
        run.font.size = Pt(9)
        run.italic = True

        # Grade entry box
        p = doc.add_paragraph()
        run = p.add_run("FINAL GRADE: _______ / " + str(int(pts)) + "    NOTES: _________________________________")
        run.font.size = Pt(11)

        # Separator
        doc.add_paragraph("_" * 80)

    # Footer summary
    doc.add_page_break()
    add_section_heading(doc, "QUICK REFERENCE -- GRADE ENTRY")

    final_table = doc.add_table(rows=1, cols=5)
    final_table.style = "Light Grid Accent 1"
    for i, h in enumerate(["#", "Student", "Assignment", "Suggested", "Final"]):
        final_table.rows[0].cells[i].text = h
        for p in final_table.rows[0].cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)

    for sub in all_submissions:
        suggested, max_pts, _ = suggest_grade(sub)
        row = final_table.add_row()
        row.cells[0].text = str(sub["entry_num"])
        row.cells[1].text = sub["student_name"]
        row.cells[2].text = sub["assignment_name"][:35]
        if suggested is not None:
            row.cells[3].text = f"{suggested:.0f}/{max_pts:.0f}"
        else:
            row.cells[3].text = "Review"
        row.cells[4].text = "____"
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    path = os.path.join(DOWNLOADS, "Grading_Reference_04-13-2026.docx")
    doc.save(path)
    print(f"Saved: {path}")
    return path


if __name__ == "__main__":
    build_grading_doc()
